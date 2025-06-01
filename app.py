import os
import json
import io
import base64
import streamlit as st
import openai
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from docx.shared import Pt
from typing import Optional, List, Dict, Tuple, Union, BinaryIO
import tempfile
from pathlib import Path
import markdown
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import frontmatter
from dotenv import load_dotenv

# Config file path
CONFIG_FILE = 'config.json'

def load_config():
    """Load configuration from file"""
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def save_config(config):
    """Save configuration to file"""
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, indent=2)

# Load saved config
config = load_config()

# Initialize session state with saved values or defaults
if 'api_key' not in st.session_state:
    st.session_state.api_key = config.get('api_key', '')
if 'api_base' not in st.session_state:
    st.session_state.api_base = config.get('api_base', 'https://api.openai.com/v1')
if 'models' not in st.session_state:
    st.session_state.models = []
if 'selected_model' not in st.session_state:
    st.session_state.selected_model = config.get('selected_model', 'gpt-4')

def extract_text_from_pdf(file) -> str:
    pdf_reader = PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file) -> str:
    doc = Document(file)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def extract_text_from_txt(file) -> str:
    return file.getvalue().decode("utf-8")

def extract_text_from_srt(file) -> str:
    return file.getvalue().decode("utf-8")

def extract_text_from_md(file) -> str:
    content = file.getvalue().decode("utf-8")
    # Extract content without frontmatter
    post = frontmatter.loads(content)
    return post.content

def create_download_link(content: str, filename: str, button_text: str) -> str:
    """Generate a download link for the given content"""
    b64 = base64.b64encode(content.encode()).decode()
    return f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">{button_text}</a>'

def save_as_pdf(text: str, filename: str) -> bytes:
    """Convert text to PDF and return as bytes"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Normal', fontName='Helvetica', fontSize=12, leading=14))
    
    # Split text into paragraphs and create story
    story = []
    for para in text.split('\n\n'):
        if para.strip():
            p = Paragraph(para, styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 12))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def save_as_docx(text: str, filename: str) -> bytes:
    """Convert text to DOCX and return as bytes"""
    doc = Document()
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Add paragraphs
    for para in text.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def split_into_chunks(text: str, max_chunk_size: int = 10000) -> List[str]:
    """Split text into chunks of specified size, trying to break at paragraph boundaries."""
    chunks = []
    while len(text) > max_chunk_size:
        # Try to find the last paragraph break within the chunk
        split_pos = text.rfind('\n\n', 0, max_chunk_size)
        # If no paragraph break found, try to split at sentence boundary
        if split_pos == -1:
            split_pos = text.rfind('. ', 0, max_chunk_size)
        # If still no good place to split, split at max_chunk_size
        if split_pos == -1 or split_pos < max_chunk_size // 2:
            split_pos = max_chunk_size
        
        chunks.append(text[:split_pos].strip())
        text = text[split_pos:].strip()
    
    if text:
        chunks.append(text)
    return chunks

def translate_text(
    text: str, 
    target_language: str, 
    model: str,
    tone: str = "neutral",
    check_grammar: bool = False
) -> Tuple[str, str]:
    try:
        client = openai.OpenAI(
            api_key=st.session_state.api_key,
            base_url=st.session_state.api_base
        )
        
        # Prepare system message based on options
        system_message = f"""You are a professional translator. 
            Translate the following text to {target_language} with a {tone} tone while preserving:
            1. All formatting and structure
            2. Technical terms and proper names (do not translate names, brands, etc.)
            3. Any special characters or codes
            4. The overall style of the original text
            
            If the text contains placeholders like {{variable}}, HTML tags, or special formatting, 
            keep them exactly as they are in the translation."""
            
        if check_grammar:
            system_message += """
            
            After translation, carefully review the text for any grammatical, punctuation, 
            or stylistic errors. For each correction, add a detailed note at the end of the 
            document in the following format:
            
            --- CORRECTIONS MADE ---
            [Original text] -> [Corrected text]
            [Explanation of the correction and why it was made]
            """
        
        # Split text into manageable chunks
        chunks = split_into_chunks(text)
        if not chunks:
            return "", ""
            
        translated_chunks = []
        corrections = []
        
        with st.spinner("Translating..."):
            progress_bar = st.progress(0)
            total_chunks = len(chunks)
            
            for i, chunk in enumerate(chunks):
                try:
                    # Update progress
                    progress = (i) / total_chunks
                    progress_bar.progress(min(progress, 1.0))
                    
                    # Prepare user message
                    user_message = f"""Language to translate to: {target_language}
                    Tone: {tone}
                    Check grammar: {'Yes' if check_grammar else 'No'}
                    
                    Text to translate:
                    {chunk}"""
                    
                    # Translate chunk
                    response = client.chat.completions.create(
                        model=model,
                        messages=[
                            {"role": "system", "content": system_message},
                            {"role": "user", "content": user_message}
                        ],
                        temperature=0.3,
                    )
                    
                    translated_text = response.choices[0].message.content
                    
                    # Extract corrections if present
                    if check_grammar and "--- CORRECTIONS MADE ---" in translated_text:
                        translated_text, correction = translated_text.split("--- CORRECTIONS MADE ---", 1)
                        corrections.append(correction.strip())
                    
                    translated_chunks.append(translated_text.strip())
                    
                except Exception as e:
                    error_msg = f"Error translating chunk {i+1}/{len(chunks)}: {str(e)}"
                    st.error(error_msg)
                    translated_chunks.append(f"[TRANSLATION ERROR: {error_msg}]")
            
            # Update progress to 100%
            progress_bar.progress(1.0)
        
        # Combine all translated chunks
        full_translation = "\n\n".join(translated_chunks)
        
        # Add corrections section if any
        full_corrections = ""
        if check_grammar and corrections:
            full_corrections = "\n\n--- CORRECTIONS MADE ---\n\n" + "\n\n".join(corrections)
        
        return full_translation, full_corrections
        
    except Exception as e:
        st.error(f"Error during translation: {str(e)}")
        return ""

def main():
    st.set_page_config(page_title="AI Document Translator Pro", layout="wide")
    st.title("AI Document Translator Pro")
    
    # Custom CSS for better UI
    st.markdown("""
    <style>
    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    .stButton>button {
        width: 100%;
    }
    .uploadedFile {
        background-color: #f0f2f6;
        border-radius: 5px;
        padding: 10px;
        margin-bottom: 10px;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Sidebar for API settings
    with st.sidebar:
        st.header("API Configuration")
        
        # API Key Input
        new_api_key = st.text_input("API Key", value=st.session_state.api_key, type="password")
        if new_api_key != st.session_state.api_key:
            st.session_state.api_key = new_api_key
            config = load_config()
            config['api_key'] = new_api_key
            save_config(config)
            st.rerun()
            
        # API Base URL Input
        new_api_base = st.text_input("API Base URL", value=st.session_state.api_base)
        if new_api_base != st.session_state.api_base:
            st.session_state.api_base = new_api_base
            config = load_config()
            config['api_base'] = new_api_base
            save_config(config)
            st.rerun()
        
        # Model selection
        if st.button("Load Available Models"):
            try:
                with st.spinner("Loading models..."):
                    client = openai.OpenAI(
                        api_key=st.session_state.api_key,
                        base_url=st.session_state.api_base.rstrip('/')
                    )
                    models = client.models.list()
                    available_models = [model.id for model in models.data]
                    
                    if available_models:
                        st.session_state.models = available_models
                        if not st.session_state.selected_model in available_models:
                            st.session_state.selected_model = available_models[0]
                            
                        # Save config
                        config = load_config()
                        config.update({
                            'api_key': st.session_state.api_key,
                            'api_base': st.session_state.api_base,
                            'selected_model': st.session_state.selected_model
                        })
                        save_config(config)
                        
                        st.success(f"Successfully loaded {len(available_models)} models")
                    else:
                        st.warning("No models found. Please check your API key and endpoint.")
                        
            except Exception as e:
                st.error(f"Failed to load models: {str(e)}")
        
        if st.session_state.models:
            # Create a selectbox for model selection
            selected = st.selectbox(
                "Select Model",
                st.session_state.models,
                index=st.session_state.models.index(st.session_state.selected_model) if st.session_state.selected_model in st.session_state.models else 0,
                key="model_selector"
            )
            
            # Save model selection when changed
            if selected != st.session_state.selected_model:
                st.session_state.selected_model = selected
                config = load_config()
                config['selected_model'] = selected
                save_config(config)
                st.rerun()
    
    # Main content
    st.write("### Upload Documents")
    uploaded_files = st.file_uploader(
        "Choose one or more files",
        type=["pdf", "docx", "txt", "srt", "md"],
        accept_multiple_files=True
    )
    
    # Translation options
    st.write("### Translation Options")
    col1, col2 = st.columns(2)
    with col1:
        tone = st.selectbox(
            "Tone of Translation",
            ["Neutral", "Formal", "Casual", "Academic", "Business", "Friendly"],
            index=0
        )
    with col2:
        check_grammar = st.checkbox(
            "Check grammar and suggest improvements",
            value=False,
            help="Review the text for grammatical errors and suggest improvements"
        )
    
    # Output format selection
    output_format = st.radio(
        "Output Format",
        ["Same as input", "PDF", "DOCX", "TXT"],
        horizontal=True
    )
    
    # Language selection
    languages = {
        "English": "English",
        "Slovenian": "Slovenian",
        "German": "German",
        "French": "French",
        "Spanish": "Spanish",
        "Italian": "Italian",
        "Croatian": "Croatian",
        "Serbian": "Serbian",
        "Russian": "Russian",
        "Chinese": "Chinese",
        "Japanese": "Japanese",
        "Korean": "Korean"
    }
    
    col1, col2 = st.columns(2)
    with col1:
        source_lang = st.selectbox("Source Language", list(languages.keys()), index=0)
    with col2:
        target_lang = st.selectbox("Target Language", list(languages.keys()), index=1)
    
    translate_clicked = st.button("Translate Documents")
    
    if translate_clicked and uploaded_files:
        if not st.session_state.api_key:
            st.error("Please enter your API key first.")
        elif not st.session_state.selected_model:
            st.error("Please load and select a model first.")
        else:
            for uploaded_file in uploaded_files:
                with st.expander(f"{uploaded_file.name}", expanded=True):
                    st.write(f"### Processing: {uploaded_file.name}")
                    
                    try:
                        # Extract text based on file type
                        file_ext = uploaded_file.name.lower().split('.')[-1]
                        if file_ext == 'pdf':
                            text = extract_text_from_pdf(uploaded_file)
                        elif file_ext == 'docx':
                            text = extract_text_from_docx(uploaded_file)
                        elif file_ext == 'txt':
                            text = extract_text_from_txt(uploaded_file)
                        elif file_ext == 'srt':
                            text = extract_text_from_srt(uploaded_file)
                        elif file_ext == 'md':
                            text = extract_text_from_md(uploaded_file)
                        else:
                            st.error(f"Unsupported file type: {file_ext}")
                            continue
                        
                        with st.spinner(f"Translating {uploaded_file.name}..."):
                            # Translate the text
                            translated_text, corrections = translate_text(
                                text=text,
                                target_language=languages[target_lang],
                                model=st.session_state.selected_model,
                                tone=tone.lower(),
                                check_grammar=check_grammar
                            )
                            
                            # Add corrections to the end of the document if any
                            if corrections:
                                translated_text += "\n\n" + corrections
                            
                            # Determine output format
                            if output_format == "Same as input":
                                out_ext = file_ext
                            else:
                                out_ext = output_format.lower()
                            
                            # Prepare file for download based on format
                            out_filename = f"translated_{Path(uploaded_file.name).stem}.{out_ext}"
                            
                            if out_ext in ['pdf', 'PDF']:
                                file_data = save_as_pdf(translated_text, out_filename)
                                mime_type = "application/pdf"
                            elif out_ext in ['docx', 'DOCX']:
                                file_data = save_as_docx(translated_text, out_filename)
                                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            else:  # Default to TXT
                                file_data = translated_text.encode('utf-8')
                                mime_type = "text/plain"
                            
                            # Display download buttons
                            st.success(f"Translation of {uploaded_file.name} completed!")
                            
                            # Show preview in a container
                            st.write("### Translation Preview")
                            with st.container(border=True):
                                st.text_area("Translated Text", 
                                           value=translated_text, 
                                           height=300, 
                                           key=f"translated_{uploaded_file.name}",
                                           label_visibility="collapsed")
                            
                            st.write("### Download Options")
                            # Download buttons
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.download_button(
                                    label=f"Download as {out_ext.upper()}",
                                    data=file_data,
                                    file_name=out_filename,
                                    mime=mime_type,
                                    key=f"dl_{uploaded_file.name}"
                                )
                            
                            # Always show TXT download option if output is not already TXT
                            if out_ext != 'txt':
                                with col2:
                                    st.download_button(
                                        label="Download as TXT",
                                        data=translated_text.encode('utf-8'),
                                        file_name=f"translated_{Path(uploaded_file.name).stem}.txt",
                                        mime="text/plain",
                                        key=f"txt_{uploaded_file.name}"
                                    )
                            
                            # Show copy to clipboard button
                            with col3:
                                st.download_button(
                                    label="Copy to Clipboard",
                                    data=translated_text,
                                    file_name="clipboard.txt",
                                    mime="text/plain",
                                    key=f"copy_{uploaded_file.name}"
                                )
                            
                    except Exception as e:
                        st.error(f"An error occurred while processing {uploaded_file.name}: {str(e)}")
                        import traceback
                        st.text(traceback.format_exc())
    elif translate_clicked and not uploaded_files:
        st.warning("Please upload at least one file to translate.")

if __name__ == "__main__":
    main()
